import redcap_parser as rp
import data2docx as d2d
from docx import Document
from docx.shared import Pt
import utilities as util
import parse as pp
if __name__ == '__main__':
    
    comps       = pp.parse('/home/guillaume/Documents/redcap/data.csv')
    all_tasks   = util.get_all_tasks(comps)

    stype=1

    summary_data = util.summary(comps,stype)
    d2d.summary_comp2docx(summary_data,stype)


    len(summary_data)

    summary_data[2]




#%% Check if have component

data_tnames         = ['T8.2.1','T8.4.1']
data_factor_tnames  = ['T8.1.1','T8.5.2',''] 

task_names = ['T8.3.11',]

has_v = util.has_tasks(summary_data,task_names,3)


has_v

#%% PUT TASKS FROM ALL COMPONENTS TOGTHER AND SAVE


print 'total number of tasks: ', len(all_tasks)
document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(6)

tree = d2d.sorted_task2docx(document,all_tasks)

document.save('redcap_sorted.docx')    
    
    
#%%

test = [['18', 'First dataset annotated according to the ontology.'], ['12', 'Prototype of the ontology for describing data on patients with neurological diseases developed'], ['24', 'Ontology and several datasets annotated according to the ontology.']]


#%% SAVE ONE COMPONENT (TESTING)

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(6)

row_id              = 2
col_names, data_val = rp.extrat_row(ws,row_id)   
comp                = util.row2comp(col_names,data_val)

print 'comp(', row_id,') has ', len(comp.tasks) ,' tasks'

comps.append(comp)

d2d.comp2docx(document,comp,bSorted=True)   
    
document.save('test4.docx')    

    
#%% SAVE ALL COMPONENTS (TESTING)

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(6)

i = 0
for comp in comps:
    print '=====> ', i, '<====='
    d2d.comp2docx(document,comp,bSorted=False)   
    i = i + 1

document.save('test3.docx')    











import redcap_parser as rp
import data2docx as d2d
from docx import Document
from docx.shared import Pt
import utilities as util

if __name__ == '__main__':
    
    filename = "/home/guillaume/Documents/redhat/data3.csv"
    data     = rp.import_csv(filename)
        
    
    row_id = 0
    col_names, data_val = rp.csv2data(data,row_id)
    comp                = util.row2comp(col_names,data_val)


#%%

# Extracting components

num_rows            = len(data)-1
comps               = []


for row_id in range(0,num_rows):
    
    col_names, data_val = rp.csv2data(data,row_id)
    comp                = util.row2comp(col_names,data_val)
    
    print 'comp(', row_id,') has ', len(comp.tasks) ,' tasks'
    comps.append(comp)
       
    
#%% Print

print ' '
for comp in comps:
    print comp    
    print ' '
    
#%% import csv

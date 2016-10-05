#!/usr/bin/env python2
# -*- coding: utf-8 -*-
"""
Created on Mon Oct  3 17:24:32 2016

@author: guillaume
"""

from docx import Document
import redcap_parser as rp
from tree import Tree


def one_line_2_docx(p,name,value=''):
    """
        Prints to word document format:
            
            <b>name</b>value\n
    
    """
    # print component    
    p.add_run(name).bold = True
    if value != '':
        p.add_run(value)
        p.add_run('\n')
        
        
        
def comp2docx(document,comp,bSorted=False):
    
    
    p = document.add_paragraph()
    
    p.add_run('----------------------------------------------------------------------------------------------------------------------------------------------------\n')   
    one_line_2_docx(p,'Component(' + comp.subject_ID  + ') \n')           
    one_line_2_docx(p,'Task leader: ', comp.task_leader      )
    
    if bSorted == False:
        for task in comp.tasks:
            print 'processing task: ', task.name 
            task2docx(document,task)
    else:
        
        comp.sort()
        leaf_nodes = comp.tree.get_leaf_nodes()
        
        
        for leaf in leaf_nodes:
        
            # PRINT HEADING OF ONE LEAF
            
            elements = leaf.path
            i = 1
            for elem in elements:
                document.add_heading(elem, level=i)
                i = i + 1
                
            p.add_run('\n')
    
            for idx in leaf.idx:
                
                task2docx(document,comp.tasks[idx],b_print_heading=False)
                
def sorted_task2docx(document,tasks):

    tree = Tree()
    for i in range(0,len(tasks)):
        if len(tasks[i].build_block_belong) == 3:
            tree.add_element(tasks[i].build_block_belong,i)   
            
            
    p = document.add_paragraph()       
            
                
    leaf_nodes = tree.get_leaf_nodes()
    for leaf in leaf_nodes:
                # PRINT HEADING OF ONE LEAF
           
            elements = leaf.path
            i = 1
            for elem in elements:
                document.add_heading(elem, level=i)
                i = i + 1
                
            p.add_run('\n')
    
            for idx in leaf.idx:
                task2docx(document,tasks[idx],b_print_heading=False)
     
        
    
def task2docx(document,task,b_print_heading=True):    
    
    # To which building block your component belongs to
    
    if b_print_heading:
        elements = task.build_block_belong # ('SOFTWARE', 'Data Factory (DF)', 'Feature Engineering')
        i = 1;
        for elem in elements:
            document.add_heading(elem, level=i)
            i = i + 1
    
    bPrint = True
        
    if bPrint == True:
        #Prior Dependencies table:
        p = document.add_paragraph()
        p.style = 'Normal'
    
        one_line_2_docx(p,'Contributing task: ',task.task_number)
        one_line_2_docx(p,'Description: '      ,task.short_desc_comp)

    
        one_line_2_docx(p,'Dependencies:')
    
        #  add table
    
        data_hosp = task.build_block_data.get('Hospital Data',[])
        data_ref  = task.build_block_data.get('Reference data',[])
    
        soft      = task.build_block_soft
        serv      = task.build_block_serv
        
    
        dependencies2docx(document,data_hosp,data_ref,soft,serv)

    
        # Prior Release table:
        p = document.add_paragraph()
        p.add_run('\n')
        one_line_2_docx(p,'Releases:')
        releases2docx(document,task.planned_functionality)
    
        # Prior User table:
        p = document.add_paragraph()
        p.add_run('\n')
        one_line_2_docx(p,'User and Use cases:')   
        usecase2docx(document,task.users,task.short_desc_usecase)    
    
            
    

def dependencies2docx(wdoc,data_hosp,data_ref,soft,services):
    """
        
    Takes Data and software information and puts it into a table
    
    Prameters
    ---------
     
    wdoc: python-docx word document object
    data_hosp: DATA > Hospital Data > [] list
    data_ref : DATA > Reference Data > [] list
    soft     : dict
    services : dict
    
    """
    
    data_hosp = ', '.join(data_hosp)
    data_ref  = ', '.join(data_ref)
    
    
    if len(data_hosp) == 0:
        bDataHosp = False
    else:
        bDataHosp = True
        
    if len(data_ref) == 0:
        bDataRef = False
    else:
        bDataRef = True
        
        
   
    # get number of sofware 
    num_soft = len(soft.keys())
    
    # get number of services
    num_serv = len(services.keys())
    
    
    tcase = 'full'
    
    if bDataHosp & bDataRef  :
        tab_num_rows = 2 + num_soft + num_serv
        tcase        = 'full'    
    elif (bDataHosp == True) & (bDataRef == False):
        tab_num_rows = 1 + num_soft + num_serv
        tcase        = 'only_hosp'         
    elif (bDataHosp == False) & (bDataRef == True):
        tab_num_rows = 1 + num_soft + num_serv
        tcase        = 'only_ref'         
    else:
        tab_num_rows = num_soft + num_serv  
        tcase        = 'only_soft'         
    
        
        
        
    if tab_num_rows == 0:
        return
            
    table       = wdoc.add_table(rows=tab_num_rows, cols=3)
    table.style = 'TableGrid'
   
    

    if tcase == 'full':
        
        # Hospital Data
        row = table.rows[0]
        row.cells[0].text = 'DATA'
        row.cells[1].text = 'Hospital'
        row.cells[2].text = data_hosp

        # Reference Data
        row = table.rows[1]
        row.cells[1].text = 'Reference'
        row.cells[2].text = data_ref

        r_idx = 2

    elif tcase == 'only_hosp':
        
        # Hospital Data
        row = table.rows[0]
        row.cells[0].text = 'DATA'
        row.cells[1].text = 'Hospital'
        row.cells[2].text = data_hosp

        r_idx = 1
        
    elif tcase == 'only_ref':

        # Reference Data
        row = table.rows[0]
        row.cells[1].text = 'Reference'
        row.cells[2].text = data_ref

        r_idx = 1
    
    else:
    
        r_idx = 0
        
    
    print tcase, r_idx, tab_num_rows
        

    if r_idx < len(table.rows):
        # Software Data
        row = table.rows[r_idx]
        row.cells[0].text = 'SOFTWARE'
    
        idx = r_idx    
        for key in soft:
            row = table.rows[idx]
            soft_desc = ', '.join(soft[key])
            row.cells[1].text = key
            row.cells[2].text = soft_desc
            idx = idx + 1
        r_idx = idx    
    else:
        
        print '[Error] r_idx >= len(table.rows):   r_idx: ', r_idx, '  len(table.rows):  ', len(table.rows) 
     
    if num_serv != 0:
        row = table.rows[r_idx]
        row.cells[0].text = 'SERVICES'   
        idx = r_idx    
        
        
        
        for key in services:
             row = table.rows[idx]
             serv_desc = ', '.join(services[key])
             row.cells[1].text = key
             row.cells[2].text = serv_desc
             idx = idx + 1
        r_idx = idx    

        
        
    

def releases2docx(wdoc,release):
    """
        Takes release information  (Planned functional etc..)
    
    
    """
    
    num_releases = len(release.keys())
       
    table = wdoc.add_table(rows=num_releases, cols=2)
    table.style = 'TableGrid'

    
    idx = 0    
    for key in release:
        row = table.rows[idx]
        row.cells[0].text = key
        row.cells[1].text = release[key]

        idx = idx + 1

def usecase2docx(wdoc,users,short_desc_usecase):
    
    
    table = wdoc.add_table(rows=2, cols=2)
    table.style = 'TableGrid'
    

    # Heading
    row = table.rows[0]
    row.cells[0].text = 'User'
    row.cells[1].text = 'Use cases'

    # User & Use cases    
    row = table.rows[1]
    row.cells[0].text = users
    row.cells[1].text = short_desc_usecase   

    
def component2docx(document,comp):
    
    
    
    print 'comp.titles: ', comp.titles
    
    for i in range(0,len(comp.titles)):
        document.add_heading(comp.titles[i], level=i+1)
        
    
    # Titles
    #document.add_heading(comp.titles[0], level=1)
    #document.add_heading(comp.titles[1], level=2)
    #document.add_heading(comp.titles[2], level=3)
    
    
    p = document.add_paragraph()
    
    one_line_2_docx(p,'Component(' + comp.subject_ID  + '): ',            comp.component)
    one_line_2_docx(p,'Contributing task(s): ', comp.contribution)
    one_line_2_docx(p,'Description: ',          comp.short_desc)
    
    
    # Prior Dependencies table:
    p = document.add_paragraph()
    p.style = 'Normal'
    
    one_line_2_docx(p,'Dependencies:')
    
    
    
    dependencies2docx(document,comp.data_hosp,comp.data_ref,comp.soft)

    # Prior Release table:
    p = document.add_paragraph()
    p.add_run('\n')
    one_line_2_docx(p,'Releases:')
    releases2docx(document,comp.release)
    
    
    # Prior User table:
    p = document.add_paragraph()
    p.add_run('\n')
    one_line_2_docx(p,'User and Use cases:')   
    usecase2docx(document,comp.use_case)
    

def data2word(data_vals,col_names,document):
    """ 
        data_vals   : list 
        
        col_names   : list
    
        document    : word document 
    
    """
    subject_ID = data_vals[0]
    document.add_heading('Subject: ' + subject_ID, level=3)
    p = document.add_paragraph('')
    run = p.add_run()
    run.add_break()    

    for i in range(1,len(col_names)):
        col_name = col_names[i]
        data_val = data_vals[i]
        

        if data_val != 'None':
            run = p.add_run()
            run.font.bold = True
            run.add_text(col_name + '\n')
            run = p.add_run()
            run.font.bold = False
            run.add_text(data_val + '\n\n')

            
            run.add_break()    
   

def save2word_document(ws):
    document = Document()

    document.add_heading('redcap_data', 0)
    document.add_paragraph('automatically generated from python')

    num_rows = ws.max_row

    for i in range(2,num_rows):
        col_names,data_val = rp.extrat_row(ws,i)
        print data_val
        data2word(data_val,col_names,document)

    document.save('redcap_python.docx')    
    return data_val, col_names
            